#!/usr/bin/env python3
"""
NYC Council Fiscal Impacts Pipeline
Fetches fiscal impact statement .docx files from NYC Legistar, extracts structured
data using the Claude API, and writes results to fiscal_impacts.json for the
data website.

Usage:
    python fetch_fiscal_impacts.py
    python fetch_fiscal_impacts.py --incremental
    python fetch_fiscal_impacts.py --incremental --historical
    python fetch_fiscal_impacts.py --incremental --historical --historical-years 2010-2023
    python fetch_fiscal_impacts.py --dry-run
    python fetch_fiscal_impacts.py --help

Environment variables required:
    ANTHROPIC_API_KEY   Your Anthropic API key

Notes:
    - Uses Legistar's public web interface (no API token required)
    - Caches downloaded .docx files in pipeline/cache/docx/ to avoid re-downloading
    - Run with --incremental in GitHub Actions to only process new matters
"""

from __future__ import annotations

import os
import re
import json
import time
import logging
import argparse
from datetime import datetime
from pathlib import Path
from typing import Optional

import requests
from docx import Document
import anthropic

# ── Logging ─────────────────────────────────────────────────────────────────
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s %(levelname)s %(message)s",
    datefmt="%H:%M:%S",
)
log = logging.getLogger(__name__)

# ── Paths ────────────────────────────────────────────────────────────────────
SCRIPT_DIR   = Path(__file__).parent
REPO_ROOT    = SCRIPT_DIR.parent   # data_website/ root (this repo)
CACHE_DIR    = SCRIPT_DIR / "cache" / "docx"
OUTPUT_PATH  = REPO_ROOT / "nyc_council_fiscal_impacts_tracker" / "data" / "fiscal_impacts.json"
BASE_URL     = "https://legistar.council.nyc.gov"
CLAUDE_MODEL = "claude-haiku-4-5-20251001"

# ── Extraction prompt ─────────────────────────────────────────────────────────
EXTRACTION_PROMPT = """You are extracting structured data from a New York City Council fiscal impact statement document. Extract the following fields and return ONLY a valid JSON object — no markdown fences, no explanation, just the JSON.

DOCUMENT TEXT:
---
{text}
---

Return a JSON object with exactly these fields (use null for missing/unknown, 0 for explicit zeros, true/false for booleans):

{{
  "file_number": "e.g. Int. No. 805  or  T2026-1631  or  Res. No. 1234-A  — as written",
  "legislation_type": "Introduction or Resolution or Pre-Considered Resolution or Other",
  "title": "full title as written",
  "committee": "committee name only (e.g. Transportation, Aging, Finance)",
  "sponsors": ["Last Name 1", "Last Name 2"],
  "prime_sponsor": "Last Name of first listed sponsor",
  "effective_date": "as written (e.g. 120 days after becoming law)",
  "fy_first_effective": "e.g. FY26 or FY27 — just the FY label",
  "fy_full_impact": "e.g. FY27 or FY28 — just the FY label",
  "source_of_funds": "General Fund or N/A or Federal Funds or as written",
  "cost_estimable": true,

  "total_revenue": 0,
  "total_expenditure": 0,
  "total_capital": null,
  "net_fiscal_impact": 0,

  "fiscal_table_columns": [
    {{
      "label": "column header exactly as written",
      "revenue": 0,
      "expenditure": 0,
      "capital": null,
      "net": 0
    }}
  ],

  "agencies_abbrev": ["DOT", "DPR"],
  "agencies_full": ["Department of Transportation", "Department of Parks and Recreation"],

  "program_breakdowns": [
    {{
      "agency": "DOT",
      "program": "program or line item name",
      "description": "one-sentence description",
      "cost_type": "expense or capital or revenue",
      "amount": 1000000,
      "fy_range": "FY22-FY26 or similar",
      "offset_notes": "any note about baseline offsets, or null"
    }}
  ],

  "impact_narrative_revenue": "full text of the Impact on Revenues paragraph",
  "impact_narrative_expenditure": "full text of the Impact on Expenditures paragraph",

  "omb_estimate_provided": false,
  "omb_estimate_notes": "what OMB said, or null if section absent",

  "estimate_prepared_by": "Name, Title",
  "estimate_reviewed_by": ["Name, Title"],
  "date_prepared": "YYYY-MM-DD if parseable, otherwise as written",
  "hearing_date": null
}}

RULES:
- total_revenue / total_expenditure / total_capital: sum across ALL fiscal-year columns. Always positive numbers.
- net_fiscal_impact = total_revenue - total_expenditure - total_capital. Negative = net cost to city.
- If ANY cell says "See below" or indicates the cost cannot be estimated, set cost_estimable to false and set that total to null.
- fiscal_table_columns must preserve the exact column structure from the document (there may be 2–6 columns).
- agencies_abbrev: list only agencies that are directly responsible for implementing the legislation — i.e. agencies that have at least one line item in program_breakdowns. Do NOT list agencies that only appear in passing in narrative text (e.g. OMB as reviewer, IBO as analyst, NYC Council as introducer).
- Standard NYC agency abbreviations: DOT, DPR, NYPD, FDNY, DOE, DSS, DFTA, DEP, HPD, HRA, DCAS, DSNY, DOF, DOB, DHS, NYCEM, TLC, SBS, DYCD, DOHMH, DCA, DDC, MTA, DOC, DCLA, ACS, MOCJ, OMB. Create reasonable abbreviations for others.
- program_breakdowns: extract named cost line items from the Impact on Expenditures section. May be empty [].
- For program_breakdowns entries involving street sign installation, street sign fabrication, co-naming of thoroughfares, or sign procurement: set agency="DOT" regardless of which agency the document credits. DOT is responsible for all street signage in NYC.
- For sponsors and prime_sponsor: strip all prefixes ("Council Member", "Council Members", "By Council Members", "(s):"). Return only the name. For "The Speaker (Council Member X)", return "X (Speaker)". Always use last name only as written in the document.
- Return ONLY the JSON object — no markdown, no explanation.

NARRATIVE FORMAT (older documents without a structured table):
Some documents — particularly pre-2019 legislation — state fiscal impacts as prose rather than a year-by-year table. If there is no structured numeric table, synthesize the totals from the narrative text using these rules:
- Look for phrases like "estimated to cost $X", "increase expenditures by $X annually", "reduce revenues by $X", "capital cost of $X", "estimated at $X million".
- If a dollar amount is given as an annual figure with no multi-year breakdown, use that figure as the total (do not multiply by years unless the document explicitly states a total cumulative cost).
- Revenue REDUCTIONS (e.g. "this legislation would reduce revenues by $204,000") are a cost to the city: set total_revenue = 0 and add the reduction amount to total_expenditure so net_fiscal_impact is negative.
- "No impact on revenues" or "existing resources" means 0 for that category — do NOT set cost_estimable to false.
- If a range is given (e.g. "$1 million to $2 million"), use the midpoint.
- Create a single fiscal_table_columns entry with label "Total" and populate revenue/expenditure/capital/net from the narrative figures.
- If the narrative gives a cost figure but says it "cannot be estimated precisely" or "will be determined", set cost_estimable to false.
"""


# ── Legistar scraping ─────────────────────────────────────────────────────────

def create_session() -> requests.Session:
    s = requests.Session()
    s.headers["User-Agent"] = (
        "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 "
        "(KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36"
    )
    return s


def search_legistar_all(session: requests.Session) -> list[tuple[str, str]]:
    """
    Search Legistar for ALL matters with 'Fiscal Impact Statement' in attachments,
    across all years, handling pagination.

    Background: Legistar's year filter (lstYears) is non-functional for attachment
    searches — it always returns the same results regardless of the year selected.
    We therefore search with 'All Years' and paginate through all result pages.
    Pagination uses ASP.NET __doPostBack with the RadGrid pager event targets.

    Returns list of (matter_id, guid) tuples, deduplicated.
    """
    log.info("Searching Legistar for all fiscal impact statement attachments ...")

    r = session.get(f"{BASE_URL}/Legislation.aspx", timeout=20)
    r.raise_for_status()

    vs_match  = re.search(r'id="__VIEWSTATE"\s+value="([^"]+)"', r.text)
    vsg_match = re.search(r'id="__VIEWSTATEGENERATOR"\s+value="([^"]+)"', r.text)
    if not vs_match:
        log.error("Could not find __VIEWSTATE on Legistar search page")
        return []

    # Initial search POST — "All Years" returns all available bills
    post_data = {
        "__VIEWSTATE":          vs_match.group(1),
        "__VIEWSTATEGENERATOR": vsg_match.group(1) if vsg_match else "",
        "ctl00$ContentPlaceHolder1$txtSearch":    "Fiscal Impact Statement",
        "ctl00$ContentPlaceHolder1$lstYears":     "All Years",
        "ctl00$ContentPlaceHolder1$lstTypeBasic": "All Types",
        "ctl00$ContentPlaceHolder1$chkAttachments": "on",
        "ctl00$ContentPlaceHolder1$btnSearch":    "Search Legislation",
    }

    r2 = session.post(f"{BASE_URL}/Legislation.aspx", data=post_data, timeout=30)
    r2.raise_for_status()

    seen:   set[str]           = set()
    unique: list[tuple[str, str]] = []

    def _extract_matters(html: str) -> None:
        for mid, guid in re.findall(
            r"LegislationDetail\.aspx\?ID=(\d+)&(?:amp;)?GUID=([A-F0-9\-]+)", html
        ):
            if mid not in seen:
                seen.add(mid)
                unique.append((mid, guid))

    _extract_matters(r2.text)
    log.info(f"  Page 1: {len(unique)} matters")

    # Paginate: find all numeric page links beyond page 1 in the pager.
    # The RadGrid pager renders links as:
    #   __doPostBack('ctl00$...$ctl04','')  → page 2
    #   __doPostBack('ctl00$...$ctl06','')  → page 3  etc.
    # We detect them by finding the pager HTML and extracting event targets
    # for pages 2, 3, … until no new pages are found.
    current_html = r2.text
    page_num = 1

    while True:
        # Find pager link for the next page (page_num + 1).
        # Pager links are <a href="javascript:__doPostBack(...)"><span>N</span></a>
        # The current page has class="rgCurrentPage" with no href navigation.
        next_page = page_num + 1
        # Match: href with doPostBack target followed by <span>{next_page}</span>
        pattern = (
            r"doPostBack\(&#39;([^&]+)&#39;,&#39;&#39;\)"
            r"[^<]*<span>" + str(next_page) + r"</span>"
        )
        m = re.search(pattern, current_html)
        if not m:
            break  # no more pages

        event_target = m.group(1)
        vs2  = re.search(r'id="__VIEWSTATE"\s+value="([^"]+)"', current_html)
        vsg2 = re.search(r'id="__VIEWSTATEGENERATOR"\s+value="([^"]+)"', current_html)

        page_data = {
            "__VIEWSTATE":          vs2.group(1) if vs2 else "",
            "__VIEWSTATEGENERATOR": vsg2.group(1) if vsg2 else "",
            "__EVENTTARGET":        event_target,
            "__EVENTARGUMENT":      "",
            "ctl00$ContentPlaceHolder1$txtSearch":    "Fiscal Impact Statement",
            "ctl00$ContentPlaceHolder1$lstYears":     "All Years",
            "ctl00$ContentPlaceHolder1$lstTypeBasic": "All Types",
            "ctl00$ContentPlaceHolder1$chkAttachments": "on",
        }
        rn = session.post(f"{BASE_URL}/Legislation.aspx", data=page_data, timeout=30)
        rn.raise_for_status()

        before = len(unique)
        _extract_matters(rn.text)
        added = len(unique) - before
        log.info(f"  Page {next_page}: {added} new matters (running total: {len(unique)})")

        current_html = rn.text
        page_num = next_page
        time.sleep(1)

    log.info(f"  -> {len(unique)} total matters found across all pages")
    return unique


def search_legistar_advanced_year(
    session: requests.Session, year: str, max_retries: int = 2
) -> list[tuple[str, str]]:
    """
    Use the Legistar advanced search form to find matters with 'Fiscal Impact Statement'
    attachments for a specific year. The advanced lstYearsAdvanced filter appears to
    work for historical data where the basic lstYears filter does not.

    Two-step process:
    1. GET the page, then POST with btnSwitch to enter advanced mode
    2. POST with txtAtt + lstYearsAdvanced to run the search, then paginate
    """
    log.info(f"  Advanced search for year {year} ...")

    for attempt in range(max_retries + 1):
        try:
            # Step 1: GET the page
            r = session.get(f"{BASE_URL}/Legislation.aspx", timeout=20)
            r.raise_for_status()

            vs_match  = re.search(r'id="__VIEWSTATE"\s+value="([^"]+)"', r.text)
            vsg_match = re.search(r'id="__VIEWSTATEGENERATOR"\s+value="([^"]+)"', r.text)
            if not vs_match:
                log.error(f"  Could not find __VIEWSTATE for year {year}")
                return []

            # Step 2: Switch to advanced search mode
            switch_data = {
                "__VIEWSTATE":          vs_match.group(1),
                "__VIEWSTATEGENERATOR": vsg_match.group(1) if vsg_match else "",
                "ctl00$ContentPlaceHolder1$btnSwitch":    "Advanced search >>>",
                "ctl00$ContentPlaceHolder1$lstYears":     "This Year",
                "ctl00$ContentPlaceHolder1$lstTypeBasic": "All Types",
            }
            r2 = session.post(f"{BASE_URL}/Legislation.aspx", data=switch_data, timeout=20)
            r2.raise_for_status()

            vs2  = re.search(r'id="__VIEWSTATE"\s+value="([^"]+)"', r2.text)
            vsg2 = re.search(r'id="__VIEWSTATEGENERATOR"\s+value="([^"]+)"', r2.text)

            # Step 3: Run the advanced search for this year
            client_state = json.dumps({
                "enabled": True,
                "emptyMessage": "",
                "validationText": year,
                "valueAsString": year,
                "lastSetTextInitiatesRequest": False,
                "blockAnimationTimer": 0,
                "lastAutoCompleteIndex": -1,
                "Direction": 0,
            })
            search_data = {
                "__VIEWSTATE":          vs2.group(1) if vs2 else "",
                "__VIEWSTATEGENERATOR": vsg2.group(1) if vsg2 else "",
                "ctl00$ContentPlaceHolder1$txtAtt":               "Fiscal Impact Statement",
                "ctl00$ContentPlaceHolder1$lstYearsAdvanced":     year,
                "ctl00_ContentPlaceHolder1_lstYearsAdvanced_ClientState": client_state,
                "ctl00$ContentPlaceHolder1$lstType":              "All Types",
                "ctl00$ContentPlaceHolder1$lstMax":               "50",
                "ctl00$ContentPlaceHolder1$btnSearch":            "Search Legislation",
            }
            r3 = session.post(f"{BASE_URL}/Legislation.aspx", data=search_data, timeout=90)
            r3.raise_for_status()

            seen:   set[str]              = set()
            unique: list[tuple[str, str]] = []

            def _extract(html: str) -> None:
                for mid, guid in re.findall(
                    r"LegislationDetail\.aspx\?ID=(\d+)&(?:amp;)?GUID=([A-F0-9\-]+)", html
                ):
                    if mid not in seen:
                        seen.add(mid)
                        unique.append((mid, guid))

            _extract(r3.text)
            log.info(f"    Year {year} page 1: {len(unique)} matters")

            # Paginate using the same RadGrid __doPostBack pattern
            current_html = r3.text
            page_num = 1

            while True:
                next_page = page_num + 1
                pattern = (
                    r"doPostBack\(&#39;([^&]+)&#39;,&#39;&#39;\)"
                    r"[^<]*<span>" + str(next_page) + r"</span>"
                )
                m = re.search(pattern, current_html)
                if not m:
                    break

                event_target = m.group(1)
                vs_p  = re.search(r'id="__VIEWSTATE"\s+value="([^"]+)"', current_html)
                vsg_p = re.search(r'id="__VIEWSTATEGENERATOR"\s+value="([^"]+)"', current_html)

                page_data = {
                    "__VIEWSTATE":          vs_p.group(1) if vs_p else "",
                    "__VIEWSTATEGENERATOR": vsg_p.group(1) if vsg_p else "",
                    "__EVENTTARGET":        event_target,
                    "__EVENTARGUMENT":      "",
                    "ctl00$ContentPlaceHolder1$txtAtt":               "Fiscal Impact Statement",
                    "ctl00$ContentPlaceHolder1$lstYearsAdvanced":     year,
                    "ctl00_ContentPlaceHolder1_lstYearsAdvanced_ClientState": client_state,
                    "ctl00$ContentPlaceHolder1$lstType":              "All Types",
                    "ctl00$ContentPlaceHolder1$lstMax":               "50",
                }
                rn = session.post(f"{BASE_URL}/Legislation.aspx", data=page_data, timeout=60)
                rn.raise_for_status()

                before = len(unique)
                _extract(rn.text)
                added = len(unique) - before
                log.info(f"    Year {year} page {next_page}: {added} new matters (total: {len(unique)})")

                current_html = rn.text
                page_num = next_page
                time.sleep(1.5)

            log.info(f"    Year {year}: {len(unique)} total matters")
            return unique

        except requests.exceptions.Timeout:
            if attempt < max_retries:
                wait = 30 * (attempt + 1)
                log.warning(f"    Timeout for year {year} (attempt {attempt+1}) — retrying in {wait}s")
                time.sleep(wait)
            else:
                log.error(f"    Timeout for year {year} after {max_retries+1} attempts — skipping")
                return []
        except Exception as e:
            log.error(f"    Error searching year {year}: {e}")
            return []

    return []


def get_fiscal_attachment(
    session: requests.Session, matter_id: str, guid: str
) -> tuple[str | None, str | None]:
    """
    Fetch a matter's detail page and find the fiscal impact attachment.
    Returns (attachment_id, attachment_guid) or (None, None).
    """
    url = (
        f"{BASE_URL}/LegislationDetail.aspx"
        f"?ID={matter_id}&GUID={guid}&Options=Attachments|&Search=Fiscal+Impact+Statement"
    )
    r = session.get(url, timeout=20)
    r.raise_for_status()

    views = re.findall(
        r"View\.ashx\?M=F&ID=(\d+)&(?:amp;)?GUID=([A-F0-9\-]+)", r.text
    )

    for att_id, att_guid in views:
        att_url = f"{BASE_URL}/View.ashx?M=F&ID={att_id}&GUID={att_guid}"
        try:
            head = session.head(att_url, timeout=10, allow_redirects=True)
            cd = head.headers.get("Content-Disposition", "")
            if "fiscal" in cd.lower() or "impact" in cd.lower():
                return att_id, att_guid
        except Exception:
            continue

    return None, None


def download_docx(
    session: requests.Session, att_id: str, att_guid: str
) -> Path | None:
    """
    Download a .docx attachment to the cache directory.
    Returns the local path, or None on failure.
    """
    CACHE_DIR.mkdir(parents=True, exist_ok=True)
    local = CACHE_DIR / f"{att_id}.docx"

    if local.exists():
        log.info(f"  Using cached docx: {att_id}.docx")
        return local

    url = f"{BASE_URL}/View.ashx?M=F&ID={att_id}&GUID={att_guid}"
    r = session.get(url, timeout=30)
    r.raise_for_status()

    if len(r.content) < 100:
        log.warning(f"  Suspiciously small download for att_id={att_id}")
        return None

    local.write_bytes(r.content)
    log.info(f"  Downloaded: {att_id}.docx ({len(r.content):,} bytes)")
    return local


def extract_docx_text(docx_path: Path) -> str:
    """Extract text from a .docx, including table cell content."""
    try:
        doc = Document(str(docx_path))
    except Exception as e:
        log.warning(f"  Could not open {docx_path}: {e}")
        return ""

    parts = []

    for para in doc.paragraphs:
        t = para.text.strip()
        if t:
            parts.append(t)

    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            row_text = "\t".join(cells)
            if any(cells):
                parts.append(row_text)

    return "\n".join(parts)


# ── Claude extraction ─────────────────────────────────────────────────────────

def extract_fiscal_data(
    text: str, client: anthropic.Anthropic
) -> dict:
    """Call Claude API to extract structured fiscal data from docx text."""
    prompt = EXTRACTION_PROMPT.format(text=text[:18000])

    for attempt in range(3):
        try:
            msg = client.messages.create(
                model=CLAUDE_MODEL,
                max_tokens=4096,
                messages=[{"role": "user", "content": prompt}],
            )
            raw = msg.content[0].text.strip()

            # Strip markdown fences if present
            if raw.startswith("```"):
                raw = re.sub(r"^```(?:json)?\n?", "", raw)
                raw = re.sub(r"\n?```$", "", raw.rstrip())

            data = json.loads(raw)
            return data

        except json.JSONDecodeError as e:
            log.warning(f"  JSON decode error (attempt {attempt+1}): {e}")
            if attempt == 2:
                return {"extraction_error": str(e)}
        except anthropic.RateLimitError:
            wait = 30 * (attempt + 1)
            log.warning(f"  Rate limited — waiting {wait}s")
            time.sleep(wait)
        except Exception as e:
            log.warning(f"  API error (attempt {attempt+1}): {e}")
            time.sleep(2 ** attempt)

    return {"extraction_error": "Failed after 3 attempts"}


# ── Data persistence ──────────────────────────────────────────────────────────

def load_existing(path: Path) -> dict:
    if path.exists():
        with open(path, encoding="utf-8") as f:
            return json.load(f)
    return {"metadata": {}, "records": []}


def save_output(path: Path, records: list) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    data = {
        "metadata": {
            "last_updated": datetime.utcnow().isoformat() + "Z",
            "total_records": len(records),
            "source": "NYC Legistar (legistar.council.nyc.gov)",
        },
        "records": records,
    }
    with open(path, "w", encoding="utf-8") as f:
        json.dump(data, f, indent=2, ensure_ascii=False)
    log.info(f"Saved {len(records)} records -> {path}")


# ── Fiscal impact filters ────────────────────────────────────────────────────

def text_is_zero_impact(text: str) -> bool:
    """
    Fast pre-check on raw docx text. Returns True if the document clearly
    shows all-zero figures and no 'See below' language, so we can skip the
    Claude API call entirely.
    """
    t = text.lower()
    # If the doc says cost cannot be estimated, let Claude decide
    if "see below" in t or "cannot estimate" in t or "unable to estimate" in t:
        return False
    # Find all dollar amounts in the text (e.g. $1,234,567 or $1.2 million or $0)
    amounts = re.findall(
        r"\$\s*([\d,]+(?:\.\d+)?)\s*(?:million|billion)?", t
    )
    non_zero = [a for a in amounts if a.replace(",", "").replace(".", "") not in ("0", "")]
    # If there are no dollar figures at all, or all are literally $0, skip
    return len(non_zero) == 0


def record_has_fiscal_impact(fiscal: dict) -> bool:
    """
    Post-extraction check. Returns True only if the bill has a non-zero,
    estimable fiscal impact worth storing.
    Excludes only bills where ALL of revenue, expenditure, capital, and net
    are zero — i.e. no fiscal impact whatsoever.
    """
    if not fiscal.get("cost_estimable", True):
        return False
    if "extraction_error" in fiscal:
        return False  # don't store failed extractions
    exp = fiscal.get("total_expenditure") or 0
    cap = fiscal.get("total_capital") or 0
    rev = fiscal.get("total_revenue") or 0
    net = fiscal.get("net_fiscal_impact") or 0
    return any(abs(v) > 0 for v in [exp, cap, rev, net])


def is_budget_modification(fiscal: dict) -> bool:
    """
    Returns True if this record is a mayoral budget modification (MN-#),
    not independent legislation. These are Charter §107(e) administrative
    approvals and should not appear in the fiscal tracker.
    """
    title = (fiscal.get("title") or "").lower()
    return bool(re.search(r"\bmn-\d+\b", title) or "modification (mn" in title)


def is_proposed_bill(fiscal: dict) -> bool:
    """
    Returns True if this record is a proposed (not yet passed) bill.
    Titles beginning with 'Proposed' indicate draft legislation that has
    not been enacted as a local law. Only final/passed bills belong in
    the tracker.
    """
    title = (fiscal.get("title") or "").strip()
    return title.lower().startswith("proposed")


_SIGN_KEYWORDS = [
    "street sign", "sign installation", "new street sign", "sign procurement",
    "co-name", "thoroughfare sign", "street co-name", "new signs",
    "sign fabricat", "signs at $", "signs for renamed", "signs for thoroughfare",
]

_DOT_FULL = "Department of Transportation"


def normalize_agency_attribution(fiscal: dict) -> dict:
    """
    Apply two post-extraction agency cleanup rules:

    1. Street sign line items → agency = DOT.
       Any program_breakdowns entry whose program or description mentions
       street sign installation, fabrication, or co-naming is credited to DOT.

    2. agencies_abbrev pruning.
       If program_breakdowns are present, keep only agencies that actually
       appear in at least one breakdown entry. This removes agencies that
       Claude listed from narrative text only (e.g. OMB as reviewer, IBO as
       analyst, NYC Council as introducer).
    """
    pbs = fiscal.get("program_breakdowns") or []
    if not pbs:
        return fiscal

    # Rule 1: assign DOT to sign line items
    for pb in pbs:
        combined = ((pb.get("program") or "") + " " + (pb.get("description") or "")).lower()
        if any(kw in combined for kw in _SIGN_KEYWORDS):
            pb["agency"] = "DOT"

    # Rule 2: rebuild agencies_abbrev from pb agencies only
    seen: list[str] = []
    seen_set: set[str] = set()
    for pb in pbs:
        a = pb.get("agency")
        if a and a not in seen_set:
            seen.append(a)
            seen_set.add(a)

    if seen:
        old_map = dict(zip(
            fiscal.get("agencies_abbrev") or [],
            fiscal.get("agencies_full") or [],
        ))
        old_map["DOT"] = _DOT_FULL
        fiscal["agencies_abbrev"] = seen
        fiscal["agencies_full"]   = [old_map.get(a, a) for a in seen]

    return fiscal


# ── Main ──────────────────────────────────────────────────────────────────────

def main() -> int:
    parser = argparse.ArgumentParser(description="NYC Council Fiscal Impacts Pipeline")
    parser.add_argument(
        "--incremental", action="store_true",
        help="Skip matters already present in fiscal_impacts.json",
    )
    parser.add_argument(
        "--dry-run", action="store_true",
        help="Process data but do not write output file",
    )
    parser.add_argument(
        "--historical", action="store_true",
        help="Also search historical years via the advanced search form (slow)",
    )
    parser.add_argument(
        "--historical-years", default="2014-2023",
        metavar="START-END",
        help="Year range for --historical mode, e.g. '2010-2023' (default: 2014-2023)",
    )
    args = parser.parse_args()

    api_key = os.environ.get("ANTHROPIC_API_KEY")
    if not api_key:
        log.error("ANTHROPIC_API_KEY environment variable is not set. See README.")
        return 1

    client  = anthropic.Anthropic(api_key=api_key)
    session = create_session()

    existing     = load_existing(OUTPUT_PATH)
    existing_ids = {str(r["matter_id"]) for r in existing.get("records", [])}
    records      = list(existing.get("records", []))

    total_new = 0

    # Collect all matters, deduplicating by matter_id.
    all_matter_ids: set[str] = set()
    matters: list[tuple[str, str]] = []

    def _add_matters(new: list[tuple[str, str]]) -> int:
        added = 0
        for mid, guid in new:
            if mid not in all_matter_ids:
                all_matter_ids.add(mid)
                matters.append((mid, guid))
                added += 1
        return added

    # Basic all-years search — covers all available bills in Legistar's
    # attachment index (currently 2024+; the year filter is non-functional here).
    basic = search_legistar_all(session)
    _add_matters(basic)
    log.info(f"Basic search: {len(basic)} results, {len(matters)} unique so far")
    time.sleep(2)

    # Optional historical search using the advanced form, which has a working
    # lstYearsAdvanced filter. Run year-by-year for 2014–2023 (or custom range).
    if args.historical:
        try:
            start_str, end_str = args.historical_years.split("-")
            hist_years = [str(y) for y in range(int(start_str), int(end_str) + 1)]
        except ValueError:
            log.error(f"Invalid --historical-years value: {args.historical_years!r} (expected START-END)")
            return 1

        log.info(f"Historical search: years {args.historical_years}")
        for year in hist_years:
            hist = search_legistar_advanced_year(session, year)
            added = _add_matters(hist)
            log.info(f"  Year {year}: {added} new unique matters (running total: {len(matters)})")
            time.sleep(3)  # be polite between year searches

    log.info(f"Total matters to process: {len(matters)}")

    for matter_id, guid in matters:
        if args.incremental and matter_id in existing_ids:
            log.info(f"  Skipping already-processed matter {matter_id}")
            continue

        log.info(f"Processing matter {matter_id} ...")

        try:
            att_id, att_guid = get_fiscal_attachment(session, matter_id, guid)
            time.sleep(0.5)

            if not att_id:
                log.info(f"  No fiscal impact attachment found — skipping")
                continue

            docx_path = download_docx(session, att_id, att_guid)
            time.sleep(0.5)

            if not docx_path:
                continue

            text = extract_docx_text(docx_path)
            if not text.strip():
                log.warning(f"  Empty text from {docx_path} — skipping")
                continue

            # Fast pre-check: skip obvious zero-impact bills before calling Claude.
            # If every dollar figure in the text is $0 and there's no "See below",
            # there's nothing worth storing.
            if text_is_zero_impact(text):
                log.info(f"  Pre-check: all-zero fiscal impact — skipping Claude call")
                existing_ids.add(matter_id)  # mark as seen so --incremental skips it
                continue

            log.info("  Calling Claude for extraction ...")
            fiscal = extract_fiscal_data(text, client)

            # Post-extraction filter: skip if no real fiscal impact or unestimable.
            if not record_has_fiscal_impact(fiscal):
                log.info(f"  Post-check: zero/unestimable fiscal impact — skipping")
                existing_ids.add(matter_id)
                continue

            # Skip budget modification resolutions (MN-#) — these are Charter
            # §107(e) administrative approvals, not independent legislation.
            if is_budget_modification(fiscal):
                log.info(f"  Budget modification (MN-#) — skipping")
                existing_ids.add(matter_id)
                continue

            # Skip proposed (not yet passed) bills — only final/enacted legislation
            # belongs in the tracker.
            if is_proposed_bill(fiscal):
                log.info(f"  Proposed bill (not yet passed) — skipping")
                existing_ids.add(matter_id)
                continue

            # Normalize agency attribution: assign DOT to street sign line items
            # and prune agencies not present in program_breakdowns.
            fiscal = normalize_agency_attribution(fiscal)

            record = {
                "matter_id":    matter_id,
                "legistar_guid": guid,
                "legistar_url": (
                    f"https://legistar.council.nyc.gov/LegislationDetail.aspx"
                    f"?ID={matter_id}&GUID={guid}"
                ),
                "attachment_id": att_id,
                "processed_at": datetime.utcnow().isoformat() + "Z",
                **fiscal,
            }

            records.append(record)
            existing_ids.add(matter_id)
            total_new += 1

            fn    = fiscal.get("file_number", "?")
            title = (fiscal.get("title") or "")[:60]
            log.info(f"  -> {fn}: {title}")

        except Exception as e:
            log.error(f"  Error on matter {matter_id}: {e}", exc_info=True)

        time.sleep(1)  # be polite to Legistar

    log.info(f"Processed {total_new} new matters (total in file: {len(records)})")

    if not args.dry_run:
        save_output(OUTPUT_PATH, records)
    else:
        log.info("--dry-run: not writing output file")

    return 0


if __name__ == "__main__":
    raise SystemExit(main())
